#!/usr/bin/env python3
"""
Universal Document to YAML Converter for AI Agent Playbooks

This module extracts text from various document formats (PDF, MD, TXT, DOCX, etc.)
and converts them to structured YAML format for use with AI agents and orchestration systems.
"""

import logging
import re
from pathlib import Path
from typing import Any, Dict, List

import pdfplumber
import yaml

from .extraction_patterns import extract_best_practices_and_examples
from .models import DocumentStructure
from .nlp_utils import extract_keywords

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("python-docx not available. DOCX support disabled.")

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)


class UniversalToYAMLConverter:
    """Converts various document formats to structured YAML format for AI agent knowledge bases."""

    SUPPORTED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx", ".html"}

    def __init__(self, file_path: str):
        self.file_path = Path(file_path)
        self.filename = self.file_path.stem
        self.extension = self.file_path.suffix.lower()
        self.yaml_path = self.file_path.with_suffix(".yaml")

        if self.extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {self.extension}. Supported: {self.SUPPORTED_EXTENSIONS}"
            )

    def detect_file_type(self) -> str:
        """Detect file type based on extension."""
        ext_to_type = {
            ".pdf": "pdf",
            ".md": "markdown",
            ".txt": "text",
            ".docx": "docx",
            ".html": "html",
        }
        return ext_to_type.get(self.extension, "unknown")

    def extract_text_from_file(self) -> str:
        """Extract text content based on file type."""
        file_type = self.detect_file_type()

        if file_type == "pdf":
            return self.extract_text_from_pdf()
        elif file_type == "markdown":
            return self.extract_text_from_markdown()
        elif file_type == "text":
            return self.extract_text_from_txt()
        elif file_type == "docx":
            return self.extract_text_from_docx()
        elif file_type == "html":
            return self.extract_text_from_html()
        else:
            logger.warning(f"Unsupported file type: {file_type}, trying as plain text")
            return self.extract_text_from_txt()

    def extract_text_from_pdf(self) -> str:
        """Extract text content from PDF file."""
        logger.info(f"Extracting text from PDF: {self.file_path}")

        text_content = []
        try:
            with pdfplumber.open(self.file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)

            full_text = "\n\n".join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from {len(pdf.pages)} pages")
            return full_text

        except Exception as e:
            logger.error(f"Error extracting text from PDF {self.file_path}: {e}")
            return ""

    def extract_text_from_markdown(self) -> str:
        """Extract text content from Markdown file."""
        logger.info(f"Reading Markdown file: {self.file_path}")

        try:
            with open(self.file_path, encoding="utf-8") as f:
                text = f.read()
            logger.info(f"Extracted {len(text)} characters from Markdown file")
            return text
        except Exception as e:
            logger.error(f"Error reading Markdown file {self.file_path}: {e}")
            return ""

    def extract_text_from_txt(self) -> str:
        """Extract text content from plain text file."""
        logger.info(f"Reading text file: {self.file_path}")

        try:
            with open(self.file_path, encoding="utf-8") as f:
                text = f.read()
            logger.info(f"Extracted {len(text)} characters from text file")
            return text
        except Exception as e:
            logger.error(f"Error reading text file {self.file_path}: {e}")
            return ""

    def extract_text_from_docx(self) -> str:
        """Extract text content from DOCX file."""
        logger.info(f"Extracting text from DOCX: {self.file_path}")

        if not DOCX_AVAILABLE:
            logger.error("python-docx library not available for DOCX processing")
            return ""

        try:
            doc = DocxDocument(self.file_path)
            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            full_text = "\n\n".join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from DOCX file")
            return full_text

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {self.file_path}: {e}")
            return ""

    def extract_text_from_html(self) -> str:
        """Extract text content from HTML file."""
        logger.info(f"Extracting text from HTML: {self.file_path}")

        try:
            import html2text

            with open(self.file_path, encoding="utf-8") as f:
                html_content = f.read()

            h = html2text.HTML2Text()
            h.ignore_links = False
            h.ignore_images = True
            h.ignore_tables = False

            text = h.handle(html_content)
            logger.info(f"Extracted {len(text)} characters from HTML file")
            return text

        except ImportError:
            logger.warning("html2text not available, falling back to basic HTML parsing")
            try:
                from bs4 import BeautifulSoup

                with open(self.file_path, encoding="utf-8") as f:
                    soup = BeautifulSoup(f, "html.parser")
                text = soup.get_text(separator="\n\n")
                logger.info(f"Extracted {len(text)} characters from HTML file (basic parsing)")
                return text
            except ImportError:
                logger.error("Neither html2text nor beautifulsoup4 available for HTML processing")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from HTML {self.file_path}: {e}")
            return ""

    def parse_text_structure(self, text: str) -> DocumentStructure:
        """Parse extracted text into hierarchical document structure."""
        logger.info("Parsing text structure with hierarchical sections")

        # Parse hierarchical sections
        sections = self._parse_hierarchical_sections(text)

        # Extract keywords using NLP
        keywords = extract_keywords(text)

        # Extract best practices and examples
        best_practices, examples = extract_best_practices_and_examples(text, sections)

        # Create analysis with confidence levels
        analysis = {
            "keywords": keywords,
            "best_practices": best_practices,
            "examples": examples
        }

        # Create document structure
        doc_structure = DocumentStructure(
            metadata={
                "title": self.filename,
                "source": str(self.file_path),
                "source_type": self.detect_file_type(),
            },
            content={
                "sections": sections,
                "raw_text": text
            },
            analysis=analysis
        )

        return doc_structure

    def _parse_hierarchical_sections(self, text: str) -> List[dict]:
        """
        Parse text into hierarchical sections with proper nesting.

        Returns a list of section dictionaries compatible with the existing format.
        """
        lines = text.split("\n")
        sections = []
        section_stack = []  # Stack for hierarchical sections
        current_content = []

        for line in lines:
            line = line.strip()
            if not line:
                if current_content:
                    current_content.append("")
                continue

            # Check for section headers (Markdown-style and other patterns)
            header_match = re.match(r'^(#{1,6})\s+(.+)', line)  # Markdown headers
            if header_match:
                # Save current content to previous section
                self._save_current_content(section_stack, current_content)
                current_content = []

                # Extract header level and title
                level = len(header_match.group(1))
                title = header_match.group(2).strip()

                # Create new section
                new_section = {
                    "id": f"section_{len(sections)}",
                    "title": title,
                    "level": level,
                    "content": [],
                    "subsections": [],
                    "children": []
                }

                # Handle hierarchy
                self._add_section_to_hierarchy(sections, section_stack, new_section, level)

            # Check for other header patterns
            elif re.match(r'^\d+\.\s+.+$', line):  # Numbered sections like "1. Introduction"
                self._save_current_content(section_stack, current_content)
                current_content = []

                level = 1  # Assume top level for numbered sections
                title = line

                new_section = {
                    "id": f"section_{len(sections)}",
                    "title": title,
                    "level": level,
                    "content": [],
                    "subsections": [],
                    "children": []
                }

                self._add_section_to_hierarchy(sections, section_stack, new_section, level)

            elif len(line) < 100 and line.isupper() and not line.endswith('.'):
                # ALL CAPS titles (shorter than 100 chars)
                self._save_current_content(section_stack, current_content)
                current_content = []

                level = 1
                title = line

                new_section = {
                    "id": f"section_{len(sections)}",
                    "title": title,
                    "level": level,
                    "content": [],
                    "subsections": [],
                    "children": []
                }

                self._add_section_to_hierarchy(sections, section_stack, new_section, level)

            else:
                # Regular content
                current_content.append(line)

        # Save final content
        self._save_current_content(section_stack, current_content)

        # Flatten hierarchy for backward compatibility (but keep structure)
        return self._flatten_sections_hierarchy(sections)

    def _add_section_to_hierarchy(self, sections: List[dict], section_stack: List[dict],
                                new_section: dict, level: int):
        """Add a section to the hierarchy based on its level."""
        # Pop sections from stack that are at the same or higher level
        while section_stack and section_stack[-1]["level"] >= level:
            section_stack.pop()

        # Add to parent or root
        if section_stack:
            parent = section_stack[-1]
            parent["subsections"].append(new_section)
        else:
            sections.append(new_section)

        # Push to stack
        section_stack.append(new_section)

    def _save_current_content(self, section_stack: List[dict], current_content: List[str]):
        """Save current content to the appropriate section."""
        if not current_content or not section_stack:
            return

        # Find the deepest section in the stack
        target_section = section_stack[-1]

        # Filter out empty lines at start/end
        content_lines = [line for line in current_content if line.strip()]
        if content_lines:
            target_section["content"].extend(content_lines)

    def _flatten_sections_hierarchy(self, sections: List[dict]) -> List[dict]:
        """Flatten hierarchical sections into a flat list for backward compatibility."""
        flat_sections = []

        def flatten_recursive(section_list: List[dict], result: List[dict]):
            for section in section_list:
                # Create a flat version of the section
                flat_section = {
                    "id": section.get("id"),
                    "title": section["title"],
                    "level": section.get("level", 1),
                    "content": section.get("content", []),
                    "subsections": section.get("subsections", []),  # Keep existing subsections
                    "children": section.get("children", [])  # Add new children field
                }
                result.append(flat_section)

                # Recursively flatten subsections
                if section.get("subsections"):
                    flatten_recursive(section["subsections"], result)

        flatten_recursive(sections, flat_sections)
        return flat_sections

    def extract_key_concepts(self, text: str) -> Dict[str, Any]:
        """Extract key concepts and patterns from the text."""
        concepts = {"keywords": [], "patterns": [], "best_practices": [], "examples": []}

        # Extract potential keywords (capitalized words/phrases)
        keywords = re.findall(r"\b[A-Z][a-zA-Z]{3,}\b", text)
        concepts["keywords"] = list(set(keywords[:50]))  # Limit to top 50 unique

        # Look for best practices patterns
        practice_patterns = [
            r"Best Practice[s]?[:\s]+(.+)",
            r"Recommendation[s]?[:\s]+(.+)",
            r"Tip[s]?[:\s]+(.+)",
            r"Do[:\s]+(.+)",
            r"Avoid[:\s]+(.+)",
        ]

        for pattern in practice_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            concepts["best_practices"].extend(matches)

        # Look for examples
        example_patterns = [
            r"Example[s]?[:\s]+(.+)",
            r"For example[:\s]+(.+)",
            r"Such as[:\s]+(.+)",
        ]

        for pattern in example_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            concepts["examples"].extend(matches)

        return concepts

    def convert_to_yaml(self) -> bool:
        """Main conversion method."""
        try:
            # Extract text from the source file
            text = self.extract_text_from_file()
            if not text:
                logger.error(f"No text extracted from {self.file_path}")
                return False

            # Parse structure (now returns DocumentStructure directly)
            doc_structure = self.parse_text_structure(text)

            # Convert to dict for YAML output
            yaml_structure = doc_structure.model_dump()

            # Convert to YAML
            yaml_content = yaml.dump(
                yaml_structure, default_flow_style=False, allow_unicode=True, sort_keys=False, indent=2
            )

            # Save YAML file
            with open(self.yaml_path, "w", encoding="utf-8") as f:
                f.write(yaml_content)

            logger.info(f"Successfully converted {self.file_path} to {self.yaml_path}")
            return True

        except Exception as e:
            logger.error(f"Error converting {self.file_path} to YAML: {e}")
            return False


def process_directory(directory: str = "new") -> None:
    """Process all supported document files in a directory."""
    dir_path = Path(directory)
    dir_path.mkdir(exist_ok=True)  # Create directory if it doesn't exist

    supported_files = []

    for ext in UniversalToYAMLConverter.SUPPORTED_EXTENSIONS:
        pattern = f"**/*{ext}"
        files = list(dir_path.glob(pattern))
        supported_files.extend(files)

    if not supported_files:
        logger.info(f"No supported files found in {directory}")
        logger.info(f"Supported formats: {UniversalToYAMLConverter.SUPPORTED_EXTENSIONS}")
        return

    logger.info(f"Found {len(supported_files)} supported files")

    successful = 0
    failed = 0

    for file_path in supported_files:
        logger.info(f"Processing: {file_path}")
        try:
            converter = UniversalToYAMLConverter(file_path)
            success = converter.convert_to_yaml()

            if success:
                logger.info(f"✓ Converted: {file_path.name}")
                successful += 1
            else:
                logger.error(f"✗ Failed to convert: {file_path.name}")
                failed += 1
        except ValueError as e:
            logger.error(f"✗ Unsupported format: {file_path.name} - {e}")
            failed += 1
        except Exception as e:
            logger.error(f"✗ Unexpected error processing {file_path.name}: {e}")
            failed += 1

    logger.info(f"Processing completed: {successful} successful, {failed} failed")
